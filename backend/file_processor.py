"""
L.U.N.A File Processor
Handles extraction of content from uploaded files:
  Images  → base64 for vision models (llava, llama3.2-vision)
  PDF     → text via pdfplumber
  Word    → text via python-docx
  Excel   → table/stats via pandas + openpyxl
  CSV     → table via pandas
  Text    → read directly (py, js, md, txt, json…)
  Video   → key-frame extraction via opencv-python
  Audio   → metadata only (transcription needs Whisper)
"""

import os
import io
import base64
import logging
from pathlib import Path

logger = logging.getLogger("LUNA.files")

# ── File-type map ──────────────────────────────────────────────────────────────
CATEGORIES = {
    "image": [".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".tiff", ".svg"],
    "pdf":   [".pdf"],
    "word":  [".docx", ".doc"],
    "excel": [".xlsx", ".xls"],
    "csv":   [".csv"],
    "text":  [".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx", ".html",
              ".css", ".json", ".xml", ".yaml", ".yml", ".ini", ".cfg",
              ".toml", ".sh", ".bat", ".c", ".cpp", ".java", ".rs", ".go",
              ".php", ".rb", ".swift", ".kt", ".sql", ".r", ".m"],
    "video": [".mp4", ".avi", ".mov", ".mkv", ".webm", ".flv", ".wmv"],
    "audio": [".mp3", ".wav", ".ogg", ".m4a", ".aac", ".flac"],
    "archive": [".zip", ".rar"],
}

VISION_MODELS = [
    "llava", "llava:7b", "llava:13b", "llava:34b",
    "llama3.2-vision", "llama3.2-vision:11b",
    "moondream", "bakllava", "minicpm-v",
]

MAX_TEXT_CHARS = 14_000   # max chars sent to model as context
MAX_PDF_PAGES  = 25


def get_category(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    for cat, exts in CATEGORIES.items():
        if ext in exts:
            return cat
    return "unknown"


def is_vision_model(model_name: str) -> bool:
    name = model_name.lower()
    return any(v in name for v in VISION_MODELS)


def process_file(filename: str, file_bytes: bytes) -> dict:
    """
    Main entry point.
    Returns:
        {
            category : str,
            filename : str,
            text     : str | None,   # extracted document text
            images   : list[str],    # base64 image strings
            summary  : str,          # one-line description
            error    : str | None,
        }
    """
    category = get_category(filename)
    result = {
        "category": category,
        "filename": filename,
        "text":     None,
        "images":   [],
        "summary":  "",
        "error":    None,
    }

    try:
        if   category == "image":    _image(file_bytes, result)
        elif category == "pdf":      _pdf(file_bytes, result)
        elif category == "word":     _word(file_bytes, result)
        elif category == "excel":    _excel(filename, file_bytes, result)
        elif category == "csv":      _csv(file_bytes, result)
        elif category == "text":     _text(filename, file_bytes, result)
        elif category == "video":    _video(filename, file_bytes, result)
        elif category == "audio":    _audio(filename, result)
        elif category == "archive":  _archive(filename, file_bytes, result)
        else:
            result["error"] = (
                f"Unsupported file type '{Path(filename).suffix}'. "
                "Supported: images, PDF, Word, Excel, CSV, text/code, video, ZIP, RAR."
            )
    except Exception as exc:
        logger.exception("File processing failed: %s", filename)
        result["error"] = str(exc)

    return result


# ── Processors ─────────────────────────────────────────────────────────────────

def _image(file_bytes: bytes, result: dict) -> None:
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    result["images"]  = [b64]
    result["summary"] = f"Image attached ({len(file_bytes) // 1024} KB)"


def _pdf(file_bytes: bytes, result: dict) -> None:
    try:
        import pdfplumber
    except ImportError:
        result["error"] = "pdfplumber not installed. Run: python -m pip install pdfplumber"
        return

    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages[:MAX_PDF_PAGES]):
            text = page.extract_text() or ""
            # Also try to extract tables
            for table in page.extract_tables() or []:
                rows = [" | ".join(str(c or "") for c in row) for row in table]
                text += "\n[Table]\n" + "\n".join(rows)
            if text.strip():
                parts.append(f"--- Page {i + 1} ---\n{text.strip()}")

    result["text"]    = "\n\n".join(parts)[:MAX_TEXT_CHARS]
    result["summary"] = (
        f"PDF — {page_count} page(s), "
        f"{len(result['text'])} chars extracted"
        + (" (truncated)" if len("\n\n".join(parts)) > MAX_TEXT_CHARS else "")
    )


def _word(file_bytes: bytes, result: dict) -> None:
    try:
        from docx import Document
    except ImportError:
        result["error"] = "python-docx not installed. Run: python -m pip install python-docx"
        return

    doc   = Document(io.BytesIO(file_bytes))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            prefix = ""
            if para.style.name.startswith("Heading"):
                level  = para.style.name.split()[-1]
                prefix = "#" * int(level) + " " if level.isdigit() else "## "
            parts.append(prefix + para.text.strip())

    for table in doc.tables:
        parts.append("[Table]")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))

    result["text"]    = "\n".join(parts)[:MAX_TEXT_CHARS]
    result["summary"] = (
        f"Word document — {len(doc.paragraphs)} paragraphs, "
        f"{len(doc.tables)} table(s)"
    )


def _excel(filename: str, file_bytes: bytes, result: dict) -> None:
    try:
        import pandas as pd
    except ImportError:
        result["error"] = "pandas/openpyxl not installed. Run: python -m pip install pandas openpyxl"
        return

    ext = Path(filename).suffix.lower()
    if ext == ".xls":
        df = pd.read_excel(io.BytesIO(file_bytes), engine="xlrd")
    else:
        df = pd.read_excel(io.BytesIO(file_bytes), engine="openpyxl")

    _df_to_result(df, filename, result)


def _csv(file_bytes: bytes, result: dict) -> None:
    try:
        import pandas as pd
    except ImportError:
        result["error"] = "pandas not installed. Run: python -m pip install pandas"
        return

    df = pd.read_csv(io.BytesIO(file_bytes))
    _df_to_result(df, "CSV", result)


def _df_to_result(df, label: str, result: dict) -> None:
    import pandas as pd
    rows, cols = df.shape
    text  = f"File: {label}\n"
    text += f"Dimensions: {rows} rows × {cols} columns\n"
    text += f"Columns: {', '.join(df.columns.astype(str).tolist())}\n\n"
    text += f"First {min(rows, 50)} rows:\n"
    text += df.head(50).to_string(index=False)

    numeric = df.select_dtypes(include="number")
    if not numeric.empty:
        text += "\n\nNumeric column statistics:\n"
        text += numeric.describe().round(3).to_string()

    # Value counts for small categorical columns
    for col in df.select_dtypes(include="object").columns[:3]:
        vc = df[col].value_counts().head(10)
        text += f"\n\nTop values — '{col}':\n" + vc.to_string()

    result["text"]    = text[:MAX_TEXT_CHARS]
    result["summary"] = f"Spreadsheet — {rows} rows, {cols} columns"


def _text(filename: str, file_bytes: bytes, result: dict) -> None:
    try:
        content = file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        result["error"] = str(e)
        return

    ext  = Path(filename).suffix.lower()
    lang = ext.lstrip(".")
    result["text"] = (
        f"File: {filename}\n"
        f"```{lang}\n"
        f"{content[:MAX_TEXT_CHARS]}\n"
        "```"
    )
    result["summary"] = (
        f"Text/code file — {len(content)} chars"
        + (" (truncated)" if len(content) > MAX_TEXT_CHARS else "")
    )


def _video(filename: str, file_bytes: bytes, result: dict) -> None:
    try:
        import cv2
        import numpy as np
        import tempfile
    except ImportError:
        result["error"] = (
            "opencv-python not installed. Run: python -m pip install opencv-python\n"
            "Without it, Luna cannot extract video frames."
        )
        return

    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        cap   = cv2.VideoCapture(tmp_path)
        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        fps   = cap.get(cv2.CAP_PROP_FPS) or 24
        dur   = total / fps

        # Extract 4 evenly-spaced key frames
        indices    = [int(total * r) for r in (0.1, 0.33, 0.66, 0.9)]
        frames_b64 = []

        for idx in indices:
            cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ok, frame = cap.read()
            if not ok:
                continue
            h, w = frame.shape[:2]
            scale = min(640 / w, 480 / h, 1.0)
            if scale < 1:
                frame = cv2.resize(frame, (int(w * scale), int(h * scale)))
            _, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 72])
            frames_b64.append(base64.b64encode(buf.tobytes()).decode("utf-8"))

        cap.release()

        result["images"]  = frames_b64
        result["text"]    = (
            f"Video: {filename}\n"
            f"Duration: {dur:.1f}s  |  FPS: {fps:.1f}  |  "
            f"Total frames: {total}\n"
            f"Key frames extracted for analysis: {len(frames_b64)}"
        )
        result["summary"] = (
            f"Video — {dur:.1f}s, {len(frames_b64)} key frames extracted"
        )
    finally:
        os.unlink(tmp_path)


def _audio(filename: str, result: dict) -> None:
    result["text"] = (
        f"Audio file attached: {filename}\n"
        "Note: Luna cannot transcribe audio without Whisper. "
        "If you want audio transcription, install openai-whisper and ask for it to be added."
    )
    result["summary"] = f"Audio file: {filename} (transcription not supported)"


# ── Archive (ZIP / RAR) ────────────────────────────────────────────────────────

TEXT_EXTS = set(
    [".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx", ".html", ".css",
     ".json", ".xml", ".yaml", ".yml", ".ini", ".cfg", ".toml", ".sh",
     ".bat", ".c", ".cpp", ".h", ".hpp", ".java", ".rs", ".go", ".php",
     ".rb", ".swift", ".kt", ".sql", ".r", ".m", ".env", ".gitignore",
     ".dockerfile", ".makefile", ".gradle", ".properties"]
)
MAX_FILE_CHARS = 3_000   # chars per extracted file
MAX_FILES_EXTRACTED = 40  # max files to extract text from


def _extract_members(members, read_fn) -> tuple[list, list, int]:
    """Given a list of member names and a read(name)->bytes function,
    return (file_list, text_blocks, skipped_count)."""
    file_list   = [m for m in members if not m.endswith("/")]
    text_blocks = []
    skipped     = 0

    for name in file_list:
        # skip common noise
        parts = Path(name).parts
        if any(p in ("__pycache__", ".git", "node_modules", ".idea") for p in parts):
            continue
        ext = Path(name).suffix.lower()
        if ext not in TEXT_EXTS:
            continue
        if len(text_blocks) >= MAX_FILES_EXTRACTED:
            skipped += 1
            continue
        try:
            raw     = read_fn(name)
            content = raw.decode("utf-8", errors="replace")
            lang    = ext.lstrip(".")
            snippet = content[:MAX_FILE_CHARS]
            trunc   = "... (truncated)" if len(content) > MAX_FILE_CHARS else ""
            text_blocks.append(f"=== {name} ===\n```{lang}\n{snippet}{trunc}\n```")
        except Exception:
            pass

    return file_list, text_blocks, skipped


def _archive(filename: str, file_bytes: bytes, result: dict) -> None:
    """Extract structure + text/code content from ZIP or RAR archives."""
    import zipfile
    ext = Path(filename).suffix.lower()

    file_list: list  = []
    text_blocks: list = []
    skipped = 0

    if ext == ".zip":
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                members = zf.namelist()
                file_list, text_blocks, skipped = _extract_members(
                    members, lambda n: zf.read(n)
                )
        except zipfile.BadZipFile as e:
            result["error"] = f"Invalid ZIP file: {e}"
            return
        except Exception as e:
            result["error"] = f"Could not read ZIP: {e}"
            return

    elif ext == ".rar":
        try:
            import rarfile
        except ImportError:
            result["error"] = (
                "rarfile library not installed.\n"
                "Run: pip install rarfile\n"
                "Also ensure WinRAR/unrar is on your PATH."
            )
            return
        try:
            with rarfile.RarFile(io.BytesIO(file_bytes)) as rf:
                members = rf.namelist()
                file_list, text_blocks, skipped = _extract_members(
                    members, lambda n: rf.read(n)
                )
        except Exception as e:
            result["error"] = f"Could not read RAR: {e}"
            return
    else:
        result["error"] = f"Unsupported archive format: {ext}"
        return

    # ── Build output ────────────────────────────────────────────────────────────
    tree  = f"Archive: {filename}\n"
    tree += f"Total files: {len(file_list)}\n"
    if skipped:
        tree += f"(Text extraction limit reached; {skipped} file(s) not shown)\n"
    tree += "\nFile listing:\n"
    tree += "\n".join(f"  {f}" for f in file_list[:150])
    if len(file_list) > 150:
        tree += f"\n  … and {len(file_list) - 150} more files"

    combined = tree
    if text_blocks:
        combined += "\n\n" + "\n\n".join(text_blocks)

    result["text"]    = combined[:MAX_TEXT_CHARS]
    result["summary"] = (
        f"Archive — {len(file_list)} files, "
        f"{len(text_blocks)} text/code file(s) extracted"
    )
